import docx
import docx2txt
import os
import pymysql as psql
from sys import argv


def is_new(text, type_division):
    symbols_1 = ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]
    symbols_2 = ["I", "V", "X", "Х", "L", "C", "D", "M"]

    symbols_other = [" ", "."]

    if text.upper() == "ЭПИЛОГ" and len(text) == 6:
        return "don't care", text

    start = -1
    if len(text) >= 5:
        if text[0:5] in type_division:    
            start = 6
        elif text[0:3] in type_division:
            start = 4
    else:
        return "", ""

    name, number, is_name_zone, start_name = "", "", False, -1


    for i in range(start, len(text)):
        if is_name_zone:
            name += text[i]

        if not text[i] in symbols_other and not is_name_zone:
            if not (text[i] in symbols_1 or text[i] in symbols_2):
                return "", ""
            number += text[i]
        else:
            is_name_zone = True

    return number, name.strip()


def create_file_name(number_value, number_part, part_name, number_chapter, 
                        chapter_name, way_to_save, file_name):

    f_name = f"{way_to_save}\\\\{file_name}"

    if number_value != 0:
        f_name = f"{f_name}_Том-{number_value}"
    if number_part != 0:
        f_name = f"{f_name}_Часть-{number_part}"
    if part_name != "":
        f_name = f"{f_name}-{part_name}"
    if chapter_name.upper() != "ЭПИЛОГ":
        f_name = f"{f_name}_Глава-{number_chapter}"

    delimiter = "-" if number_chapter != "end" else "_"

    if chapter_name != "":
        f_name = f"{f_name}{delimiter}{chapter_name}"

    f_name = f"{f_name}.docx"

    print_division_data(number_value, number_part, part_name, number_chapter,
                                chapter_name, f_name)

    return f_name


def clear_from_prohibited(text):
    word = ""
    for w in text:
        if not w in ["\"", "?", "<", ">", "@"]:
            word += w

    return word


def print_division_data(number_value, number_part, part_name, number_chapter, 
                                                    chapter_name, file_name):

    print_line = "-f " + file_name + " -c " + str(number_chapter)

    optioanal_params = [[" -cn ", chapter_name], [" -p ", number_part], 
                        [" -v ", number_value], [" -pn ", part_name]]

    for p in optioanal_params:
        if (str(p[1]) != "" and p[1] != 0):
            print_line += p[0] + str(p[1])
    
    print(print_line, end="\n")


def get_many_parts(way_to_file, file_name, way_to_record):
    number_value = 0

    docx.Document(f"{way_to_file}\\\\{file_name}").save(rf".\{file_name}")

    paths = []
    folder = os.getcwd()
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith('docx') and not file.startswith('~'):
                paths.append(os.path.join(root, file))

    for path in paths:
        splitted = os.path.split(path)
        folders = [os.path.splitext(splitted[1])[0]]
        while splitted[0] != f"{splitted[0][0]}:\\":
            splitted = os.path.split(splitted[0])
            folders.insert(0, splitted[1])

        images_path = os.path.join('images', *folders)
        os.makedirs(images_path, exist_ok=True)

        doc = docx.Document(path)
        docx2txt.process(path, images_path)

        rels = {}
        for rel in doc.part.rels.values():
            if isinstance(rel._target, docx.parts.image.ImagePart):
                rels[rel.rId] = os.path.basename(rel._target.partname)

        count_chapter = count_part = count_value = count_general_chapter = 0
        first_chapter = part_name = chapter_name = ""
        chapter_started = use_part = use_value = False
        doc_i = docx.Document()

        for paragraph in doc.paragraphs:

            new_part = is_new(paragraph.text.strip(), ["ЧАСТЬ", "Часть", "часть"])
            new_chapter = is_new(paragraph.text.strip(), ["ГЛАВА", "Глава", "глава"])
            new_value = is_new(paragraph.text.strip(), ["ТОМ", "Том", "том"])

            if len(new_value[0]) > 0 or len(new_part[0]) > 0 or (len(new_chapter[0]) > 0 
                and not use_part and not use_value):

                # save
                if chapter_started:
                    f_name = create_file_name(count_value, count_part, part_name, count_chapter, chapter_name,
                                              way_to_record, file_name)

                    out_name = clear_from_prohibited(f_name)
                    doc_i.save(out_name)

                    doc_i = docx.Document()
            
            if len(new_value[0]) > 0:
                count_value += 1
                use_value = True

            if len(new_part[0]) > 0 and new_chapter[1].upper() != "ЭПИЛОГ":
                count_part += 1
                use_part = True
                part_name = new_part[1]

            if len(new_chapter[0]) > 0:
                use_part = use_value = False

            if chapter_started and new_chapter[0].split() == first_chapter:
                count_chapter = 0

            if not chapter_started and len(new_chapter[0]) > 0:
                chapter_started = True
                first_chapter = new_chapter[0].split()

            if len(new_chapter[0]) > 0:
                count_chapter += 1
                count_general_chapter += 1
                chapter_name = new_chapter[1]

            doc_i.add_paragraph(paragraph.text)

            if 'Graphic' in paragraph._p.xml:
                for rId in rels:
                    if rId in paragraph._p.xml:
                        doc_i.add_picture(os.path.join(images_path, rels[rId]))

        f_name = create_file_name(count_value, count_part, part_name, count_chapter, chapter_name, 
                                    way_to_record, file_name)

        out_name = clear_from_prohibited(f_name)
        doc_i.save(out_name)

        # os.remove(rf".\{file_name}.docx")
        break


get_many_parts(argv[1], argv[2], argv[3])


